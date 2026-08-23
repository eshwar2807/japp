"""Turn an existing resume into a master profile.

Typing a full career history into a form is the least pleasant part of setting
this up, and the information already exists in a document. This reads that
document and proposes a profile.

Two deliberate limits:

* **Nothing is saved without review.** Extraction is a best effort over a PDF's
  text layer, and a misread date or a mangled employer would otherwise become a
  "verified fact" that the whole no-fabrication design then treats as gospel.
  The result is loaded into the profile form for you to correct, not committed.
* **Legal answers are never extracted.** Work authorisation, sponsorship,
  clearance and EEO answers do not appear on a resume, so any value for them
  would be invented. They stay blank for you to fill in.
"""

from __future__ import annotations

import logging
from typing import Any

from pydantic import BaseModel, ConfigDict, Field

from config import settings

log = logging.getLogger(__name__)

MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_TEXT_CHARS = 60_000
#: Below this the document almost certainly has no text layer - a scan or an
#: image-only export - and extraction would silently return nothing useful.
MIN_TEXT_CHARS = 200

SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "txt",
}
SUPPORTED_EXTENSIONS = {".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "txt"}


class ResumeImportError(Exception):
    """The document could not be read."""


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------


def detect_kind(filename: str, content_type: str = "") -> str:
    from pathlib import PurePath

    suffix = PurePath(filename or "").suffix.lower()
    if suffix in SUPPORTED_EXTENSIONS:
        return SUPPORTED_EXTENSIONS[suffix]
    if content_type in SUPPORTED_TYPES:
        return SUPPORTED_TYPES[content_type]
    label = suffix or content_type or "unknown"
    raise ResumeImportError(
        f"Unsupported file type '{label}'. Upload a PDF, DOCX, TXT or Markdown resume."
    )


def extract_text(data: bytes, kind: str) -> str:
    """Pull the text layer out of a resume document."""
    if len(data) > MAX_UPLOAD_BYTES:
        raise ResumeImportError(
            f"File is {len(data) / 1_048_576:.1f}MB; the limit is "
            f"{MAX_UPLOAD_BYTES // 1_048_576}MB."
        )
    if not data:
        raise ResumeImportError("The file is empty.")

    if kind == "pdf":
        text = _extract_pdf(data)
    elif kind == "docx":
        text = _extract_docx(data)
    else:
        text = data.decode("utf-8", errors="replace")

    text = "\n".join(line.rstrip() for line in text.splitlines() if line.strip())
    if len(text) < MIN_TEXT_CHARS:
        raise ResumeImportError(
            "Almost no text could be read from that file. If it is a scan or an "
            "image-only PDF, export a text version and try again."
        )
    return text[:MAX_TEXT_CHARS]


def _extract_pdf(data: bytes) -> str:
    import io

    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ResumeImportError("PDF support is not installed.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:
        raise ResumeImportError(f"Could not read that PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    import io

    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ResumeImportError("DOCX support is not installed.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ResumeImportError(f"Could not read that DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:          # skills are often laid out in tables
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# --------------------------------------------------------------------------
# Structured extraction
# --------------------------------------------------------------------------


class ImportedExperience(BaseModel):
    model_config = ConfigDict(extra="ignore")

    company: str
    title: str
    location: str = ""
    start_date: str = Field(default="", description="YYYY-MM as printed on the resume.")
    end_date: str = Field(default="", description="YYYY-MM, or empty if current.")
    is_current: bool = False
    bullets: list[str] = Field(default_factory=list)
    tech_used: list[str] = Field(default_factory=list)


class ImportedEducation(BaseModel):
    model_config = ConfigDict(extra="ignore")

    institution: str
    degree: str = ""
    field_of_study: str = ""
    location: str = ""
    end_date: str = ""


class ImportedCertification(BaseModel):
    model_config = ConfigDict(extra="ignore")

    name: str
    issuer: str = ""
    issue_date: str = ""


class ImportedProfile(BaseModel):
    """What can honestly be read off a resume. No legal or EEO fields."""

    model_config = ConfigDict(extra="ignore")

    full_name: str = ""
    email: str = ""
    phone: str = ""
    city: str = ""
    state: str = ""
    country: str = ""
    linkedin: str = ""
    github: str = ""
    portfolio: str = ""
    summary: str = ""
    target_titles: list[str] = Field(default_factory=list)
    skills_hard: list[str] = Field(default_factory=list)
    skills_tooling: list[str] = Field(default_factory=list)
    skills_soft: list[str] = Field(default_factory=list)
    experience: list[ImportedExperience] = Field(default_factory=list)
    education: list[ImportedEducation] = Field(default_factory=list)
    certifications: list[ImportedCertification] = Field(default_factory=list)
    uncertain: list[str] = Field(
        default_factory=list,
        description="Anything ambiguous or hard to read, so it can be checked.",
    )


IMPORT_SYSTEM = """\
You transcribe a resume into structured fields. You are transcribing, not \
writing.

Rules:
- Copy what the document says. Do not improve, embellish, summarise or infer.
- Never invent an employer, title, date, degree, certification, metric or tool. \
If the resume does not say it, leave the field empty.
- Dates: normalise to YYYY-MM. "Jan 2021" -> "2021-01". A year alone -> \
"YYYY-01", and note it in `uncertain`. "Present"/"Current" -> is_current true \
and end_date empty.
- Bullets: copy each accomplishment as written, minus the bullet character. \
Keep every metric exactly as printed.
- Classify skills: `skills_hard` for competencies (e.g. "distributed systems"), \
`skills_tooling` for named products and languages (e.g. "Kubernetes", "Go"), \
`skills_soft` for interpersonal ones. Only skills the resume actually lists.
- `summary` is the resume's own summary if it has one, otherwise empty. Do not \
compose one.
- `target_titles`: roles this person has held or is clearly aiming at.
- PDF text extraction garbles things. Anything you had to guess at - a run-on \
date, a merged column, an unclear employer boundary - goes in `uncertain` so a \
human checks it.
- Do not attempt work authorisation, sponsorship, clearance, salary or \
demographic fields. They are not on a resume and are handled elsewhere."""


class ResumeImporter:
    def __init__(self, client: Any | None = None, model: str = "",
                 api_key: str | None = None, on_usage: Any = None) -> None:
        # Transcription, not judgement: the cheap tier is the right fit.
        self.model = model or settings.LLM_MODEL_BULK
        self.api_key = api_key
        self.on_usage = on_usage
        self._client = client

    @property
    def client(self) -> Any:
        if self._client is None:
            import anthropic

            key = self.api_key or settings.ANTHROPIC_API_KEY
            self._client = anthropic.Anthropic(**({"api_key": key} if key else {})).messages
        return self._client

    def parse(self, text: str) -> ImportedProfile:
        from engine.llm import request_params

        response = self.client.parse(
            model=self.model,
            max_tokens=settings.LLM_MAX_TOKENS,
            system=IMPORT_SYSTEM,
            messages=[{"role": "user", "content":
                       "<RESUME>\n" + text + "\n</RESUME>\n\nTranscribe it."}],
            output_format=ImportedProfile,
            **request_params(self.model),
        )
        if self.on_usage is not None:
            try:
                self.on_usage("resume_import", response)
            except Exception:
                log.exception("Usage callback failed during resume import")

        parsed = getattr(response, "parsed_output", None)
        if parsed is None:
            raise ResumeImportError("Could not extract a profile from that resume.")
        return parsed


# --------------------------------------------------------------------------
# Merge into the profile schema
# --------------------------------------------------------------------------


def to_profile(imported: ImportedProfile, existing: dict | None = None) -> dict:
    """Build a profile draft, preserving anything already filled in.

    Existing values win: a resume must not overwrite an answer the user typed
    themselves, and legal fields are carried across untouched because a resume
    never contained them in the first place.
    """
    from web.profile_form import blank_profile

    profile = dict(existing) if existing else blank_profile()
    profile.setdefault("contact", {}).setdefault("location", {})
    profile["contact"].setdefault("links", {})

    def keep(current: Any, incoming: Any) -> Any:
        return current if current else incoming

    contact = profile["contact"]
    contact["full_name"] = keep(contact.get("full_name"), imported.full_name)
    contact["email"] = keep(contact.get("email"), imported.email)
    contact["phone"] = keep(contact.get("phone"), imported.phone)
    contact["location"]["city"] = keep(contact["location"].get("city"), imported.city)
    contact["location"]["state"] = keep(contact["location"].get("state"), imported.state)
    contact["location"]["country"] = keep(
        contact["location"].get("country"), imported.country or "United States")
    for key, value in (("linkedin", imported.linkedin), ("github", imported.github),
                       ("portfolio", imported.portfolio)):
        contact["links"][key] = keep(contact["links"].get(key), value)

    profile["summary"] = keep(profile.get("summary"), imported.summary)
    profile["target_titles"] = profile.get("target_titles") or imported.target_titles

    skills = profile.setdefault("skills", {})
    skills["hard"] = skills.get("hard") or imported.skills_hard
    skills["tooling"] = skills.get("tooling") or imported.skills_tooling
    skills["soft"] = skills.get("soft") or imported.skills_soft

    if not profile.get("experience"):
        profile["experience"] = [{
            "company": e.company, "title": e.title, "location": e.location,
            "start_date": e.start_date,
            "end_date": None if e.is_current else (e.end_date or None),
            "is_current": e.is_current, "employment_type": "Full-time",
            "bullets": e.bullets, "tech_used": e.tech_used,
        } for e in imported.experience]

    if not profile.get("education"):
        profile["education"] = [{
            "institution": e.institution, "degree": e.degree,
            "field_of_study": e.field_of_study, "location": e.location,
            "start_date": "", "end_date": e.end_date, "gpa": None,
        } for e in imported.education]

    if not profile.get("certifications"):
        profile["certifications"] = [{
            "name": c.name, "issuer": c.issuer, "issue_date": c.issue_date,
            "expiry_date": None, "credential_id": "",
        } for c in imported.certifications]

    # Legal and EEO answers are never touched: a resume does not contain them.
    profile.setdefault("legal", {})
    profile.setdefault("voluntary_disclosures", {})
    profile.setdefault("preferences", {})
    return profile
